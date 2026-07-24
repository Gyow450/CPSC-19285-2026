from functools import wraps
from io import BytesIO
import os
import sqlite3
from typing import IO
from werkzeug.security import generate_password_hash, check_password_hash

from docx import Document
from docx.document import Document as DocumentObject
from docx.shared import Pt
from flask import Flask, render_template, request, jsonify, session, redirect, url_for, send_file

from src.data import CPSC_Data, PipCalculate

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "dev-secret-key-change-in-production")

# 默认管理员账号：admin / admin123
# 首次运行后建议立即修改默认密码
DB_PATH = "users.db"

def init_db()->None:
    """初始化数据库，创建用户表并插入默认管理员"""
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            is_admin INTEGER DEFAULT 0
        )
    """)
    # 插入默认管理员（首次运行后应删除或改密码）
    admin_hash = generate_password_hash("admin123")
    conn.execute("""
        INSERT OR IGNORE INTO users (id, username, password_hash, is_admin)
        VALUES (1, 'admin', ?, 1)
    """, (admin_hash,))
    conn.commit()
    conn.close()

def verify_user(username, password)->bool:
    """验证用户名密码"""
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute(
        "SELECT password_hash FROM users WHERE username = ?",
        (username,)
    ).fetchone()
    conn.close()
    
    if row is None:
        return False
    return check_password_hash(row[0], password)


def get_user(username: str) -> dict | None:
    """查询用户信息（包含管理员标识）"""
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute(
        "SELECT id, username, is_admin FROM users WHERE username = ?",
        (username,)
    ).fetchone()
    conn.close()
    if row is None:
        return None
    return {"id": row[0], "username": row[1], "is_admin": bool(row[2])}


def admin_required(view):
    """仅管理员可访问"""
    @wraps(view)
    def wrapped_view(*args, **kwargs):
        if not session.get("logged_in") or not session.get("is_admin"):
            return redirect(url_for("index"))
        return view(*args, **kwargs)
    return wrapped_view


def login_required(view):
    """未登录用户重定向到登录页，并记录原目标地址"""
    @wraps(view)
    def wrapped_view(*args, **kwargs):
        if not session.get("logged_in"):
            return redirect(url_for("login", next=request.url))
        return view(*args, **kwargs)
    return wrapped_view


def _get_float(form, key)-> float|None:
    val = form.get(key, "").strip()
    return float(val) if val else None

def _get_str(form, key)-> str|None:
    val = form.get(key, "").strip()
    return val if val else None

@app.route("/")
@login_required
def index():
    return render_template("index.html", active_page="index")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()
        if verify_user(username, password):
            user = get_user(username)
            session["logged_in"] = True
            session["username"] = username
            session["is_admin"] = user.get("is_admin", False) if user else False

            # 登录成功后跳回登录前请求的页面（仅允许同域地址）
            next_url = request.form.get("next") or request.args.get("next")
            if next_url:
                from urllib.parse import urlparse
                if urlparse(next_url).netloc == urlparse(request.url_root).netloc:
                    return redirect(next_url)
            return redirect(url_for("index"))
        return render_template("login.html", error="用户名或密码错误", next=request.args.get("next"), message=request.args.get("message"))
    return render_template("login.html", error=None, next=request.args.get("next"), message=request.args.get("message"))

@app.route("/admin/users")
@admin_required
def user_list():
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute("SELECT id, username, is_admin FROM users").fetchall()
    conn.close()
    return render_template("users.html", users=rows, active_page="users")


@app.route("/admin/users/add", methods=["POST"])
@admin_required
def user_add():
    username = request.form.get("username", "").strip()
    password = request.form.get("password", "").strip()
    is_admin = 1 if request.form.get("is_admin") == "1" else 0

    if not username or not password:
        return render_template("users.html", error="用户名和密码不能为空", users=get_all_users(), active_page="users")

    conn = sqlite3.connect(DB_PATH)
    try:
        conn.execute(
            "INSERT INTO users (username, password_hash, is_admin) VALUES (?, ?, ?)",
            (username, generate_password_hash(password), is_admin)
        )
        conn.commit()
    except sqlite3.IntegrityError:
        return render_template("users.html", error="用户名已存在", users=get_all_users(), active_page="users")
    finally:
        conn.close()
    return redirect(url_for("user_list"))


@app.route("/admin/users/<int:user_id>/reset", methods=["POST"])
@admin_required
def user_reset(user_id: int):
    new_password = request.form.get("new_password", "").strip()
    if not new_password:
        return render_template("users.html", error="新密码不能为空", users=get_all_users(), active_page="users")

    conn = sqlite3.connect(DB_PATH)
    conn.execute(
        "UPDATE users SET password_hash = ? WHERE id = ?",
        (generate_password_hash(new_password), user_id)
    )
    conn.commit()
    conn.close()
    return redirect(url_for("user_list"))


@app.route("/admin/users/<int:user_id>/delete", methods=["POST"])
@admin_required
def user_delete(user_id: int):
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute("SELECT username FROM users WHERE id = ?", (user_id,)).fetchone()
    if row and row[0] == session.get("username"):
        conn.close()
        return render_template("users.html", error="不能删除当前登录用户", users=get_all_users(), active_page="users")
    conn.execute("DELETE FROM users WHERE id = ?", (user_id,))
    conn.commit()
    conn.close()
    return redirect(url_for("user_list"))


def get_all_users():
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute("SELECT id, username, is_admin FROM users").fetchall()
    conn.close()
    return rows


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.route("/change_password", methods=["GET", "POST"])
@login_required
def change_password():
    """当前登录用户修改自己的密码"""
    username = session.get("username")
    error = None
    success = None

    if request.method == "POST":
        old_password = request.form.get("old_password", "").strip()
        new_password = request.form.get("new_password", "").strip()
        confirm_password = request.form.get("confirm_password", "").strip()

        if not new_password:
            error = "新密码不能为空"
        elif new_password != confirm_password:
            error = "两次输入的新密码不一致"
        elif not verify_user(username, old_password):
            error = "原密码错误"
        else:
            conn = sqlite3.connect(DB_PATH)
            conn.execute(
                "UPDATE users SET password_hash = ? WHERE username = ?",
                (generate_password_hash(new_password), username)
            )
            conn.commit()
            conn.close()

            # 普通用户修改本人密码后自动退出并跳回登录页
            if not session.get("is_admin"):
                session.clear()
                return redirect(url_for("login", message="密码已修改，请使用新密码重新登录"))

            success = "密码修改成功，下次登录请使用新密码"

    return render_template(
        "change_password.html",
        error=error,
        success=success,
        active_page="change_password"
    )


@app.route("/calculate", methods=["POST"])
@login_required
def calculate():
    try:
        # 前端字段 → 后端 CPSC_Data 字段映射
        params = {
            "pip_d": _get_float(request.form, "pip_d"),
            "c_type": _get_str(request.form, "coat_type"),
            "c_rg": _get_float(request.form, "rg_value"),
            "c_p": _get_float(request.form, "p_value"),
            "c_y": _get_float(request.form, "y_value"),
            "cp_exist": _get_str(request.form, "cp_exist"),
            "cp_value": _get_float(request.form, "cp_rate"),   # 注意：后面要 /100
            "soil_n": _get_float(request.form, "soil_n"),
            "soil_rho": _get_float(request.form, "soil_rho"),
            "p_move_ir": _get_float(request.form, "stable_stray_ir"),
            "p_move_noir": _get_float(request.form, "stable_stray_noir"),
            "dc_stray": _get_float(request.form, "dc_stray"),   # 注意：后面要 /100
            "ac_stray": _get_float(request.form, "ac_stray"),
            "drainage": _get_str(request.form, "drainage"),
        }

        # 保护率：前端输入的是 %（如 100），后端需要小数（1.0）
        if params["cp_value"]:
            params["cp_value"] = params["cp_value"] / 100.0
        if params["dc_stray"]:
            params["dc_stray"] = params["dc_stray"] / 100.0

        data_obj = CPSC_Data(**params)
        calc = PipCalculate(data_obj)
        matrix, matrix_latex, score, a_vec, _  = calc.calculate()
        a_text = f"结果向量A = [ {', '.join(f'{x:.4f}' for x in a_vec)} ]"
        return jsonify({
            "success": True,
            "score": round(float(score), 2),
            "matrix": matrix.tolist(),  # 5×4 嵌套数组
            "a_text": a_text,  # 1×4 数组
            "matrix_latex": matrix_latex  # LaTeX 格式的矩阵
        })

    except ValueError as e:
        return jsonify({"success": False, "error": str(e)}), 400
    except Exception as e:
        return jsonify({"success": False, "error": f"计算异常: {str(e)}"}), 500


_LABELS = {
    "pip_d": "管径（mm）",
    "c_type": "防腐层类型",
    "c_rg": "防腐层绝缘电阻率Rg值（kΩ·m²）",
    "c_p": "防腐层破损点密度P值（处/100m）",
    "c_y": "防腐层电流衰减率Y值（dB/m）",
    "cp_exist": "是否建设有阴极保护",
    "cp_value": "阴极保护率",
    "soil_n": "土壤腐蚀性评价N值",
    "soil_rho": "土壤电阻率（Ω·m）",
    "p_move_ir": "含IR降的电位正向偏移（mV）",
    "p_move_noir": "无IR降的电位正向偏移（mV）",
    "dc_stray": "阴保管道电位正于要求的比例（无阴保管道正于自然电位20mV的比例）%",
    "ac_stray": "交流电流密度（A/m²）",
    "drainage": "排流效果",
}

_MATRIX_HEADERS = [
    "外防腐层状况",
    "阴极保护有效性",
    "土壤腐蚀性",
    "杂散电流干扰",
    "排流效果",
]


def _build_docx(params: dict, score: float, matrix: list, a_text: str) -> DocumentObject:
    """根据输入参数和评价结果生成 Word 报告。"""
    doc = Document()

    title = doc.add_heading("埋地钢质管道腐蚀防护系统质量等级评价报告", level=0)
    title.alignment = 1  # 居中

    doc.add_heading("一、输入参数", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "项目"
    hdr[1].text = "数值"
    for key, label in _LABELS.items():
        value = params.get(key)
        if value is None:
            continue
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    doc.add_heading("二、评价结果", level=1)
    if score >= 90:
        level = "1"
        desc = "系统功能完好，满足设计要求，在6年的检验周期内能有效使用。"
    elif score >= 80:
        level = "2"
        desc = "系统基本完好但存在一些不影响防护效果的缺陷，能基本满足设计要求，3年~6年的检验周期内能使用。"
    elif score >= 70:
        level = "3"
        desc = "系统整体状况较差，存在缺陷，不能完全满足设计要求，在使用单位采取适当措施后，可在1年~3年检验周期内在限定的条件下使用。"
    else:
        level = "4"
        desc = "系统缺陷严重，不能满足设计要求，不能有效防止金属管体腐蚀，使用单位应采取重大维修。"

    doc.add_paragraph(f"腐蚀防护系统质量评价得分：{score:.2f}")
    doc.add_paragraph(f"等级评价：{level} 级")
    doc.add_paragraph(f"评价说明：{desc}")
    doc.add_paragraph(a_text)

    doc.add_heading("三、隶属矩阵", level=1)
    mtable = doc.add_table(rows=1, cols=5)
    mtable.style = "Light Grid Accent 1"
    headers = ["评价指标", "等级1", "等级2", "等级3", "等级4"]
    for i, h in enumerate(headers):
        mtable.rows[0].cells[i].text = h
    for i, row in enumerate(matrix):
        cells = mtable.add_row().cells
        cells[0].text = _MATRIX_HEADERS[i]
        for j, val in enumerate(row):
            cells[j + 1].text = f"{val:.3f}"

    return doc


@app.route("/download_docx", methods=["POST"])
@login_required
def download_docx():
    """根据表单数据生成评价报告 DOCX 并提供下载。"""
    try:
        params = {
            "pip_d": _get_float(request.form, "pip_d"),
            "c_type": _get_str(request.form, "coat_type"),
            "c_rg": _get_float(request.form, "rg_value"),
            "c_p": _get_float(request.form, "p_value"),
            "c_y": _get_float(request.form, "y_value"),
            "cp_exist": _get_str(request.form, "cp_exist"),
            "cp_value": _get_float(request.form, "cp_rate"),
            "soil_n": _get_float(request.form, "soil_n"),
            "soil_rho": _get_float(request.form, "soil_rho"),
            "p_move_ir": _get_float(request.form, "stable_stray_ir"),
            "p_move_noir": _get_float(request.form, "stable_stray_noir"),
            "dc_stray": _get_float(request.form, "dc_stray"),
            "ac_stray": _get_float(request.form, "ac_stray"),
            "drainage": _get_str(request.form, "drainage"),
        }

        if params["cp_value"]:
            params["cp_value"] = params["cp_value"] / 100.0
        if params["dc_stray"]:
            params["dc_stray"] = params["dc_stray"] / 100.0

        data_obj = CPSC_Data(**params)
        calc = PipCalculate(data_obj)
        matrix, _, score, a_vec, _ = calc.calculate()
        a_text = f"结果向量A = [ {', '.join(f'{x:.4f}' for x in a_vec)} ]"

        doc = _build_docx(params, float(score), matrix.tolist(), a_text)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="evaluation_report.docx",
        )
    except ValueError as e:
        return jsonify({"success": False, "error": str(e)}), 400
    except Exception as e:
        return jsonify({"success": False, "error": f"生成报告异常: {str(e)}"}), 500


if __name__ == "__main__":
    # init_db()  # 初始化数据库
    # print(f"DB_PATH: {os.path.abspath(DB_PATH)}")
    app.run(debug=True, port=5000)